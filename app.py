import streamlit as st
from supabase import create_client, Client
from docx import Document
import unicodedata
import re
import os
import io
from PIL import Image
from io import BytesIO

import streamlit.components.v1 as components

# ==========================
# CONFIGURAÇÕES DO SUPABASE
# ==========================
SUPABASE_URL = "https://fnabfzhmyrzfeqqyzeip.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImZuYWJmemhteXJ6ZmVxcXl6ZWlwIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc1Njg0MTg5OCwiZXhwIjoyMDcyNDE3ODk4fQ.0lX8Zm22DXtURKNTjKwGWvoG2N6_kuMklqvQVQsVRMk"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

ADMIN_EMAIL = "salomaopaulinomachaieie@gmail.com"

CATEGORIES = ["Medicina Geral", "Saúde Materno Infantil", "Enfermagem Geral", "Nutrição"]

# ==========================
# FUNÇÕES AUXILIARES
# ==========================

def sanitize_filename(filename: str) -> str:
    nfkd = unicodedata.normalize("NFKD", filename)
    filename = "".join([c for c in nfkd if not unicodedata.combining(c)])
    filename = filename.replace(" ", "_")
    filename = re.sub(r"[^a-zA-Z0-9_.-]", "", filename)
    return filename

def guess_mime(path: str, fallback: str = "application/octet-stream") -> str:
    ext = os.path.splitext(path.lower())[1]
    return {
        ".pdf": "application/pdf",
        ".epub": "application/epub+zip",
        ".txt": "text/plain",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(ext, fallback)

def login_user(email, password):
    try:
        user = supabase.auth.sign_in_with_password({"email": email, "password": password})
        return user
    except Exception:
        return None

def logout_user():
    st.session_state.user = None
    st.session_state.offset = 0
    st.rerun()

def get_books(search_query="", category=None, offset=0, limit=8):
    q = supabase.table("livros").select("*").order("titulo")
    if search_query:
        q = q.ilike("titulo", f"%{search_query}%")
    if category:
        q = q.eq("categoria", category)
    resp = q.range(offset, offset + limit - 1).execute()
    data = getattr(resp, "data", None)
    if data is None and isinstance(resp, dict):
        data = resp.get("data", [])
    return data or []

def record_history(user_email, livro_id):
    try:
        supabase.table("usuarios_livros_historico").insert({
            "usuario_email": user_email,
            "livro_id": livro_id
        }).execute()
    except Exception as e:
        st.error(f"Erro ao gravar histórico: {e}")



from PIL import Image
import io

def upload_book(titulo, autor, categoria, file, capa=None, largura=400, altura=600):
    try:
        # PDF
        safe_pdf = sanitize_filename(file.name)
        pdf_path = f"pdfs/{safe_pdf}"
        pdf_bytes = file.read()
        pdf_options = {"content-type": guess_mime(pdf_path, "application/pdf"), "upsert": "true"}
        supabase.storage.from_("biblioteca").upload(pdf_path, pdf_bytes, pdf_options)
        pdf_public = supabase.storage.from_("biblioteca").get_public_url(pdf_path)
        file_url = pdf_public.get("publicURL") if isinstance(pdf_public, dict) else str(pdf_public)

        # Capa opcional
        capa_url = None
        if capa:
            safe_capa = sanitize_filename(capa.name)
            capa_path = f"capas/{safe_capa}"

            img = Image.open(capa)
            img = img.resize((largura, altura), Image.Resampling.LANCZOS)

            # 👇 garantir que está em RGB (evita erro com PNG RGBA)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")

            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format="JPEG")
            capa_bytes = img_byte_arr.getvalue()

            capa_options = {"content-type": "image/jpeg", "upsert": "true"}
            supabase.storage.from_("biblioteca").upload(capa_path, capa_bytes, capa_options)
            capa_public = supabase.storage.from_("biblioteca").get_public_url(capa_path)
            capa_url = capa_public.get("publicURL") if isinstance(capa_public, dict) else str(capa_public)

        # Inserir na DB
        payload = {
            "titulo": titulo,
            "autor": autor or "Autor desconhecido",
            "categoria": categoria,
            "ficheiro_url": file_url,
            "capa_url": capa_url
        }
        supabase.table("livros").insert(payload).execute()
        return True

    except Exception as e:
        st.error(f"Erro no upload: {e}")
        return False


# ==========================
# CONFIGURAÇÃO DA PÁGINA
# ==========================

st.set_page_config(
    page_title="Biblioteca Virtual do Instituto Politécnico Sumayya",
    page_icon="static/image/logo.png",  # Caminho para o ficheiro
    layout="wide"
)

st.markdown("""
<style>
    body { background-color: #f5f7fa; }
    .stButton button { background-color: #4a90e2; color: white; border-radius: 8px; padding: 0.5em 1.5em; }
    .book-card { background: white; border-radius: 12px; padding: 15px; margin: 10px;
        box-shadow: 0 4px 10px rgba(0,0,0,0.1); transition: transform 0.2s; text-align: center; }
    .book-card:hover { transform: scale(1.03); box-shadow: 0 6px 14px rgba(0,0,0,0.15); }
    .book-cover { width: 100%; height: 250px; object-fit: cover; border-radius: 10px; margin-bottom: 10px; }
</style>
""", unsafe_allow_html=True)

# ==========================
# LOGIN
# ==========================

if "user" not in st.session_state:
    st.session_state.user = None
if "offset" not in st.session_state:
    st.session_state.offset = 0

if not st.session_state.user:
    # Centralizar logo usando colunas
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.image("static/image/logo_capa1.png", width=500, output_format="PNG", caption="")  # caminho relativo
    st.text("Seja bem-vindo ao nosso Sistema Bibliotecario! Por favor, faça login para continuar.")
    

    email = st.text_input("Email")
    password = st.text_input("Palavra-passe", type="password")
    if st.button("Entrar"):
        user = login_user(email, password)
        if user:
            st.session_state.user = {"id": user.user.id, "email": user.user.email}
            st.success(f"Bem-vindo, {email}")
            st.rerun()
        else:
            st.error("Credenciais inválidas")
    
    # Rodapé
    st.markdown(
        """
        <div style="text-align: center; margin-top: 40px; color: #888; font-size: 0.9em;">
            © 2025 Biblioteca Virtual | Fortaleza Digital, E.I | Desenvolvedor: Salomão Machaieie. Todos os direitos reservados.
        </div>
        """,
        unsafe_allow_html=True
    )
    st.stop()

    


user_email = st.session_state.user["email"]

# ==========================
# BARRA LATERAL
# ==========================
st.sidebar.markdown("## 👤 Perfil do Usuário")

# Inicializa foto do usuário se não existir
if "user_photo" not in st.session_state:
    st.session_state.user_photo = "https://cdn-icons-png.flaticon.com/512/1048/1048953.png"


# Mostrar foto atual
st.sidebar.image(st.session_state.user_photo, width=80)

# Informações do usuário
st.sidebar.markdown(f"### 👤 Utilizador: {user_email}")
if user_email == ADMIN_EMAIL:
    st.sidebar.success("Administrador")
else:
    st.sidebar.info("Usuário normal")

# Botão de logout
if st.sidebar.button("🚪 Logout"):
    logout_user()

# Filtro de categoria
CATEGORIES = ["Medicina geral", "Saúde Materno infantil", "Enfermagem geral", "Nutrição"]
category_filter = st.sidebar.selectbox("Filtrar por Categoria", ["Todas"] + CATEGORIES)

# Permitir upload de nova foto
new_photo = st.sidebar.file_uploader("Alterar foto de perfil", type=["png", "jpg", "jpeg"])
if new_photo:
    # Salva no session_state
    st.session_state.user_photo = new_photo
    st.sidebar.success("Foto atualizada!")
    st.stop()  # Interrompe execução para recarregar a sidebar com a nova imagem


# ============Carrossel========
# ============Carrossel========
# Função para obter URLs públicas de todas as imagens do bucket carrossel
def get_carousel_images():
    try:
        response = supabase.storage.from_("biblioteca").list(path="carrossel")
        files = response if isinstance(response, list) else []
        urls = []
        for file in files:
            if file.get("name").lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
                public_url = supabase.storage.from_("biblioteca").get_public_url(f"carrossel/{file['name']}")
                url = public_url.get("publicURL") if isinstance(public_url, dict) else str(public_url)
                urls.append(url)
        return urls
    except Exception as e:
        st.error(f"Erro ao obter imagens do carrossel: {e}")
        return []

carousel_images = get_carousel_images()

# Gerar HTML do carrossel dinamicamente
if carousel_images:
    carousel_items = ""
    for i, img_url in enumerate(carousel_images):
        active_class = "active" if i == 0 else ""
        carousel_items += f"""
        <div class="carousel-item {active_class}">
            <img src="{img_url}" class="d-block w-100 rounded" alt="Imagem {i+1}">
        </div>
        """

    carousel_code = f"""
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css">
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/js/bootstrap.bundle.min.js"></script>

    <style>
    .carousel-inner img {{
        height: 300px;
        object-fit: cover;
    }}
    </style>

    <div id="carouselExample" class="carousel slide mb-4" data-bs-ride="carousel" data-bs-interval="3000">
      <div class="carousel-inner">
        {carousel_items}
      </div>
      <button class="carousel-control-prev" type="button" data-bs-target="#carouselExample" data-bs-slide="prev">
        <span class="carousel-control-prev-icon"></span>
        <span class="visually-hidden">Anterior</span>
      </button>
      <button class="carousel-control-next" type="button" data-bs-target="#carouselExample" data-bs-slide="next">
        <span class="carousel-control-next-icon"></span>
        <span class="visually-hidden">Próximo</span>
      </button>
    </div>
    """

    components.html(carousel_code, height=300)
else:
    st.info("Nenhuma imagem encontrada no carrossel.")



#============= Carrossel =======


# Pesquisa
search = st.text_input("🔍 Procurar livro por título")

# ==========================
# LISTAGEM DE LIVROS (lazy load)
# ==========================
BOOKS_PER_BLOCK = 8
livros = get_books(search, None if category_filter=="Todas" else category_filter, st.session_state.offset, BOOKS_PER_BLOCK)

if livros:
    cols_per_row = 4
    for i in range(0, len(livros), cols_per_row):
        row = st.columns(cols_per_row)
        for col, livro in zip(row, livros[i:i+cols_per_row]):
            with col:
                st.markdown("<div class='book-card'>", unsafe_allow_html=True)
                capa_url = livro.get("capa_url") or "https://cdn-icons-png.flaticon.com/512/29/29302.png"
                st.image(capa_url, use_container_width=True, caption=livro["titulo"])
                st.write(f"{livro.get('autor', 'Autor desconhecido')}")
                st.markdown(f"[Baixar livro]({livro['ficheiro_url']})")
                # Grava histórico
                if user_email != ADMIN_EMAIL:
                    record_history(user_email, livro["id"])
                st.markdown("</div>", unsafe_allow_html=True)

    # Botão para carregar mais
    if len(livros) == BOOKS_PER_BLOCK:
        if st.button("Carregar mais livros"):
            st.session_state.offset += BOOKS_PER_BLOCK
            st.experimental_rerun()
else:
    st.info("Nenhum livro encontrado.")

st.markdown(
        """
        <div style="text-align: center; margin-top: 40px; color: #888; font-size: 0.9em;">
            © 2025 Biblioteca Virtual | Fortaleza Digital, E.I | Desenvolvedor: Salomão Machaieie. Todos os direitos reservados.
        </div>
        """,
        unsafe_allow_html=True
    )



# ==========================
# PAINEL ADMIN
# ==========================
if user_email == ADMIN_EMAIL:
    st.subheader("⚙️ Painel de Administração")
    with st.form("upload_form"):
        titulo = st.text_input("Título do Livro")
        autor = st.text_input("Autor")
        categoria = st.selectbox("Categoria", CATEGORIES)
        file = st.file_uploader("Upload do PDF", type=["pdf"])
        capa = st.file_uploader("Upload da Capa (opcional)", type=["png","jpg","jpeg"])
        
        # Novos campos: largura e altura
        largura = st.number_input("📏 Largura da capa (px)", min_value=100, max_value=2000, value=400, step=10)
        altura = st.number_input("📐 Altura da capa (px)", min_value=100, max_value=3000, value=600, step=10)
        
        submitted = st.form_submit_button("Enviar")
        if submitted:
            if titulo and file:
                if upload_book(titulo, autor, categoria, file, capa, largura, altura):  # 👈 passa dimensões
                    st.success("📚 Livro adicionado com sucesso!")
                    st.session_state.offset = 0
                    st.rerun()
            else:
                st.error("Preencha pelo menos o título e o PDF")

    
   # ---- Exportação de tabela para DOCX ----
    st.subheader("📄 Exportar Credenciais de Acesso")
    
    # Seleção da tabela (pode adicionar mais tabelas futuramente)
    tabela_selecionada = st.selectbox("Selecionar tabela", ["estudantes_users"])
    
    if st.button("Gerar DOCX"):
        from docx import Document
        from docx.shared import Inches, Pt
        from io import BytesIO

        try:
            # Obter dados da tabela selecionada
            response = supabase.table(tabela_selecionada).select("*").execute()
            dados = response.data
            if not dados:
                st.warning("Nenhum dado encontrado na tabela.")
            else:
                # Criar documento na memória
                doc = Document()
                doc.add_heading(f"Tabela: {tabela_selecionada}", 0)
                
                # Criar tabela no DOCX
                colunas = list(dados[0].keys())
                tabela = doc.add_table(rows=1, cols=len(colunas))
                tabela.style = 'Light List Accent 1'  # estilo mais profissional
                
                # Cabeçalho
                hdr_cells = tabela.rows[0].cells
                for i, coluna in enumerate(colunas):
                    hdr_cells[i].text = coluna
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
                
                # Preencher dados
                for item in dados:
                    row_cells = tabela.add_row().cells
                    for i, coluna in enumerate(colunas):
                        row_cells[i].text = str(item[coluna])
                        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                
                # Salvar DOCX em memória
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Botão de download
                st.download_button(
                    label="Exportar arquivo",
                    data=buffer,
                    file_name=f"{tabela_selecionada}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Ocorreu um erro ao gerar DOCX: {e}")


























